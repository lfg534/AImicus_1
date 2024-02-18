# -*- coding:utf-8 -*-
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

pl_type = {} # 判定，个人 = 0，公司 = 1

pl_name = {} # 个人
pl_race = {}
pl_id = {}
pl_birth = {}
pl_gender = {}
pl_home = {}
pl_phone = {}

pl_policy_name = {} # 保险公司
# pl_credit_num = {}
pl_policy_add = {}
pl_agent_name = {}
pl_agent_phone = {}

'''pl_islawyer = {} # 诉讼代理人信息
pl_lawyer_office = {}
pl_lawyer = {}'''

de_type = {} # 判定，个人 = 0，公司 = 1

de_name = {} # 个人
de_race = {}
de_id = {}
de_birth = {}
de_gender = {}
de_home = {}
de_phone = {}

de_policy_name = {} # 保险公司
# de_credit_num = {}
de_policy_add = {}
de_agent_name = {}
de_agent_phone = {}

'''de_islawyer = {} # 诉讼代理人信息
de_lawyer_office = {}
de_lawyer = {}'''

def set_format(paragraph):
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Pt(28)
    for run in paragraph.runs:
        run.font.name = '仿宋'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

def add_title(paragraph, title):
    run = paragraph.add_run(title)
    run.font.size = Pt(18)
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.line_spacing = 2
    for run in paragraph.runs:
        run.font.name = '宋体'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_section_title(paragraph, title):
    run = paragraph.add_run(title)
    run.font.size = Pt(14)
    run.bold = True
    # 注意这里没有设置首行缩进
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        run.font.name = '仿宋'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

def add_court(paragraph, title):
    run = paragraph.add_run(title)
    run.font.size = Pt(14)
    # 注意这里没有设置首行缩进
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        run.font.name = '仿宋'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

def add_text(paragraph, text):
    run = paragraph.add_run(text)
    run.font.size = Pt(14)
    set_format(paragraph)

def add_mixed_style_text(paragraph, texts):
    for text, is_bold in texts:
        run = paragraph.add_run(text)
        run.font.size = Pt(14)
        run.bold = is_bold
        r = run._element
    set_format(paragraph)

def add_multiple_text(doc, text):
    paragraphs = text.split('\n')
    for p in paragraphs:
        para = doc.add_paragraph()
        add_text(para, p.strip())

def add_right_aligned_text(paragraph, text):
    run = paragraph.add_run(text)
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    set_format(paragraph)


def extract_info_from_id(pl_id):
    try:
        birth_str = pl_id[6:14]
        gender_digit = int(pl_id[16])
        year = birth_str[0:4]
        month = str(int(birth_str[4:6]))
        day = str(int(birth_str[6:8]))
        birth = f"{year}年{month}月{day}日"
        # print(birth)
        if gender_digit % 2 == 0:
            gender = "女"
        else:
            gender = "男"
        return birth, gender
    
    except Exception as e:
        # print(f"Error in parsing ID: {e}")
        return "____年__月__日", "________（男/女）"


def getdocx(info_path,CHAT_ID):
    doc = Document()

    title_paragraph = doc.add_paragraph()
    add_title(title_paragraph, '民事起诉状')

    i = 0

    # Read data from the JSON file
    with open(info_path, 'r', encoding="utf-8") as f:
        data = json.load(f)

        pl_num = data.get("pl_num", 1)
        for i in range(pl_num):
            pl_type[i] = data.get(f"pl_type_{i+1}", 0)
            if pl_type[i] == 0:
                pl_name[i] = data.get(f"pl_name_{i+1}", "________")
                pl_race[i] = data.get(f"pl_race_{i+1}", "________")
                pl_id[i] = data.get(f"pl_id_{i+1}", "________")
                pl_birth[i], pl_gender[i] = extract_info_from_id(pl_id[i])
                pl_home[i] = data.get(f"pl_home_{i+1}", "________")
                pl_phone[i] = data.get(f"pl_phone_{i+1}", "________")
            elif pl_type[i] == 1:
                pl_policy_name[i] = data.get(f"pl_policy_name_{i+1}", "________")
                # pl_credit_num[i] = data.get(f"pl_credit_num_{i+1}", "________")
                pl_policy_add[i] = data.get(f"pl_policy_add_{i+1}", "________")
                pl_agent_name[i] = data.get(f"pl_agent_name_{i+1}", "________")
                pl_agent_phone[i] = data.get(f"pl_agent_phone_{i+1}", "________")
            else:
                print(f"pl_{i+1} is wrong.")
                exit()

        de_num = data.get("de_num", 1)
        for i in range(de_num):
            de_type[i] = data.get(f"de_type_{i+1}", 0)
            if de_type[i] == 0:
                de_name[i] = data.get(f"de_name_{i+1}", "________")
                de_race[i] = data.get(f"de_race_{i+1}", "________")
                de_id[i] = data.get(f"de_id_{i+1}", "________")
                de_birth[i], de_gender[i] = extract_info_from_id(de_id[i])
                de_home[i] = data.get(f"de_home_{i+1}", "________")
                de_phone[i] = data.get(f"de_phone_{i+1}", "________")
            elif de_type[i] == 1:
                de_policy_name[i] = data.get(f"de_policy_name_{i+1}", "________")
                # de_credit_num[i] = data.get(f"de_credit_num_{i+1}", "________")
                de_policy_add[i] = data.get(f"de_policy_add_{i+1}", "________")
                de_agent_name[i] = data.get(f"de_agent_name_{i+1}", "________")
                de_agent_phone[i] = data.get(f"de_agent_phone_{i+1}", "________")
            else:
                print(f"de_{i+1} is wrong.")
                exit()

        requests = data.get("requests", "请提供诉讼请求。")

        case = data.get("case", "请提供事实与经过。")

        court = data.get("court", "________人民法院")

    if pl_num == 1:
        if pl_type[0] == 0:
            plaintiff_paragraph = doc.add_paragraph()
            add_mixed_style_text(plaintiff_paragraph, \
                [(f"原告：", True), \
                (f"{pl_name[0]}, {pl_gender[0]}, {pl_race[0]}，{pl_birth[0]}出生，住{pl_home[0]}，身份证号码：{pl_id[0]}，联系电话：{pl_phone[0]}。", False)])
        else:
            plaintiff_paragraph = doc.add_paragraph()
            add_mixed_style_text(plaintiff_paragraph, \
                [(f"原告：", True), \
                (f"{pl_policy_name[0]}, 住所地：{pl_policy_add[0]}。", False)])
            plaintiff_paragraph = doc.add_paragraph()
            add_mixed_style_text(plaintiff_paragraph, \
                [(f"负责人：{pl_agent_name[0]}，联系电话：{pl_agent_phone[0]}。", False)])

    else:
        for i in range(pl_num):
            if pl_type[i] == 0:
                plaintiff_paragraph = doc.add_paragraph()
                add_mixed_style_text(plaintiff_paragraph, \
                    [(f"原告{i+1}：", True), \
                    (f"{pl_name[i]}, {pl_gender[i]}, {pl_race[i]}，{pl_birth[i]}出生，住{pl_home[i]}，身份证号码：{pl_id[i]}，联系电话：{pl_phone[i]}。", False)])
            else:
                plaintiff_paragraph = doc.add_paragraph()
                add_mixed_style_text(plaintiff_paragraph, \
                    [(f"原告{i+1}：", True), \
                    (f"{pl_policy_name[i]}, 住所地：{pl_policy_add[i]}。\n\t\t负责人：{pl_agent_name[i]}，联系电话：{pl_agent_phone[i]}。", False)])
                plaintiff_paragraph = doc.add_paragraph()
                add_mixed_style_text(plaintiff_paragraph, \
                    [(f"负责人：{pl_agent_name[i]}，联系电话：{pl_agent_phone[i]}。", False)])

    if de_num == 1:
        if de_type[0] == 0:
            defendant_paragraph = doc.add_paragraph()
            add_mixed_style_text(defendant_paragraph, \
                [("被告：", True), \
                (f"{de_name[0]}, {de_gender[0]}, {de_race[0]}，{de_birth[0]}出生，住{de_home[0]}，身份证号码：{de_id[0]}，联系电话：{de_phone[0]}。", False)])
        else:
            defendant_paragraph = doc.add_paragraph()
            add_mixed_style_text(defendant_paragraph, \
                [(f"被告：", True), \
                (f"{de_policy_name[0]}, 住所地：{de_policy_add[0]}。", False)])
            defendant_paragraph = doc.add_paragraph()
            add_mixed_style_text(defendant_paragraph, \
                [(f"负责人：{de_agent_name[0]}，联系电话：{de_agent_phone[0]}。", False)])
    
    else:
        for i in range(de_num):
            if de_type[i] == 0:
                defendant_paragraph = doc.add_paragraph()
                add_mixed_style_text(defendant_paragraph, \
                    [(f"被告{i+1}：", True), \
                    (f"{de_name[i]}, {de_gender[i]}, {de_race[i]}，{de_birth[i]}出生，住{de_home[i]}，身份证号码：{de_id[i]}，联系电话：{de_phone[i]}。", False)])
            else:
                defendant_paragraph = doc.add_paragraph()
                add_mixed_style_text(defendant_paragraph, \
                    [(f"被告{i+1}：", True), \
                    (f"{de_policy_name[i]}, 住所地：{de_policy_add[i]}。", False)])
                defendant_paragraph = doc.add_paragraph()
                add_mixed_style_text(defendant_paragraph, \
                    [(f"负责人：{de_agent_name[i]}，联系电话：{de_agent_phone[i]}。", False)])
    
    request_title_paragraph = doc.add_paragraph()
    add_section_title(request_title_paragraph, '一、诉讼请求')
    
    # request_paragraph = doc.add_paragraph()
    add_multiple_text(doc, f'{requests}')

    case_title_paragraph = doc.add_paragraph()
    add_section_title(case_title_paragraph, '二、事实与理由')

    # case_paragraph = doc.add_paragraph()
    add_multiple_text(doc, f'{case}')

    so_paragraph = doc.add_paragraph()
    add_text(so_paragraph, f'综上，为维护原告的合法权益，遂根据《中华人民共和国民法典》等相关规定，特向贵院提起诉讼，请求支持原告的诉讼请求，维护原告的合法权益。')

    so_paragraph = doc.add_paragraph()
    add_text(so_paragraph, f'此致')

    so_paragraph = doc.add_paragraph()
    add_court(so_paragraph, f'{court}')

    signature_paragraph = doc.add_paragraph()
    add_right_aligned_text(signature_paragraph, f'具状人：\t\t\t\t')

    signature_paragraph = doc.add_paragraph()
    add_right_aligned_text(signature_paragraph, f'二〇二三 年    月    日\t')


    doc.save(f'./output/起诉状_{CHAT_ID}.docx')

    return True

'''
if __name__ == '__main__':
    info_path = 'infos.json'
    getdocx(info_path)
    print(f"done")
'''